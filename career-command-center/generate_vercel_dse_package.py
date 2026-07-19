from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
import os
import subprocess
import shutil

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path("/Users/rorysemeah/Documents/RedLantern Studios")
ROOT = BASE / "rory-semeah-portfolio" / "career-command-center"
OUTPUT = ROOT / "output" / "vercel-developer-success-engineer"
RENDER_ROOT = Path("/private/tmp") / "vercel-dse-renders"
TODAY = date(2026, 7, 19)

INK = RGBColor(10, 12, 16)
SOFT = RGBColor(86, 92, 102)
LINE = RGBColor(210, 214, 220)
ACCENT = RGBColor(20, 20, 20)
PALE = RGBColor(246, 247, 249)
EMOJI_ICON_FONT = "/System/Library/Fonts/Apple Color Emoji.ttc"
EMOJI_ICON_CACHE = OUTPUT / "_emoji_icons"


@dataclass(frozen=True)
class RoleData:
    title: str
    company: str
    city: str


ROLE = RoleData(
    title="Developer Success Engineer",
    company="Vercel",
    city="San Diego, CA",
)


def set_font(run, name="Arial", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.color.rgb = INK
        style.font.size = Pt(size)


def set_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_bottom_border(paragraph, color="111111", size="8"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_run(paragraph, text, size=10.5, bold=False, italic=False, color=INK, name="Arial"):
    run = paragraph.add_run(text)
    set_font(run, name=name, size=size, bold=bold, italic=italic, color=color)
    return run


def add_para(doc, text="", *, size=10.5, bold=False, italic=False, color=INK, before=0, after=5, line=1.12, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing(p, before, after, line)
    if text:
        add_run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=10 if level == 1 else 7, after=5 if level == 1 else 3, line=1.0)
    r = p.add_run(text)
    set_font(r, size=14 if level == 1 else 11.5, bold=True, color=INK)
    return p


def add_divider(doc, color="202020"):
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=8)
    set_bottom_border(p, color=color, size="10")
    return p


def add_bullets(doc, bullets, *, size=10.25, level_prefix=""):
    for bullet in bullets:
        p = doc.add_paragraph(style="Normal")
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.08
        add_run(p, "• ", size=size, color=ACCENT, bold=True)
        add_run(p, bullet, size=size, color=INK)


def add_emoji_bullets(doc, bullets, *, size=10.25):
    for emoji, text in bullets:
        p = doc.add_paragraph(style="Normal")
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.08
        add_run(p, "• ", size=size, color=ACCENT, bold=True)
        icon_path = render_emoji_icon(emoji)
        icon_run = p.add_run()
        icon_run.add_picture(str(icon_path), width=Inches(0.15))
        add_run(p, " ", size=size, color=INK)
        add_run(p, text, size=size, color=INK)


def render_emoji_icon(emoji):
    EMOJI_ICON_CACHE.mkdir(parents=True, exist_ok=True)
    codepoints = "-".join(f"{ord(ch):x}" for ch in emoji)
    path = EMOJI_ICON_CACHE / f"{codepoints}.png"
    if path.exists():
        return path

    img = Image.new("RGBA", (128, 128), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype(EMOJI_ICON_FONT, 96)
    draw.text((0, 0), emoji, font=font, fill=(0, 0, 0, 255))
    img.save(path)
    return path


def add_contact_line(doc, parts):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=6, line=1.0)
    for idx, part in enumerate(parts):
        if idx:
            add_run(p, "  |  ", size=9.5, color=SOFT)
        add_run(p, part, size=9.75, color=SOFT, bold=False)
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=0, after=2, line=1.0)
    add_run(p, text.upper(), size=8.5, bold=True, color=SOFT)
    return p


def add_job_title_block(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=0, after=2, line=1.0)
    add_run(p, ROLE.title.upper(), size=17.5, bold=True, color=INK)
    if subtitle:
        add_run(p, "\n", size=1, color=INK)
        add_run(p, subtitle, size=9.5, color=SOFT, italic=True)
    return p


def add_credential_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=0, after=6, line=1.0)
    add_run(p, text, size=9.5, color=SOFT)
    return p


def build_resume():
    doc = Document()
    set_doc_defaults(doc)
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(fp, before=0, after=0, line=1.0)
    add_run(fp, "Rory Semeah | Developer Success Engineer", size=8.5, color=SOFT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=0, after=1, line=1.0)
    add_run(p, "Rory Semeah", size=22, bold=True, color=INK)
    add_run(p, "\n", size=1, color=INK)
    add_run(p, "Developer Success Engineer | Vercel", size=10.8, bold=True, color=ACCENT)
    add_contact_line(
        doc,
        [
            ROLE.city,
            "442-461-3093",
            "roryleesemeah@icloud.com",
            "linkedin.com/in/rory-semeah-30874555",
        ],
    )
    add_divider(doc)
    add_heading(doc, "Snapshot")
    add_emoji_bullets(
        doc,
        [
            ("🚀", "I help teams ship production web products with Next.js, Vercel, Supabase, and AI tooling."),
            ("🛠️", "I calm messy launches, troubleshoot issues, and keep delivery moving when things get noisy."),
            ("🤝", "I work across product, engineering, and operations without making the handoff heavier than it needs to be."),
            ("📚", "I write docs and playbooks so the next person does not have to ask the same question twice."),
        ],
        size=9.85,
    )
    add_heading(doc, "Selected Skills")
    add_bullets(
        doc,
        [
            "Next.js, Vercel, Supabase, OpenAI, Anthropic, AWS, Kubernetes, Terraform, REST APIs, GraphQL, CI/CD.",
            "Developer success, troubleshooting, documentation, playbooks, sprint planning, roadmap alignment, and cross functional delivery.",
        ],
        size=9.8,
    )
    add_heading(doc, "Experience")
    roles = [
        (
            "Technical Product Manager | RedLantern Studios",
            "Jan 2025 Present",
            [
                "Shipped a production multi tenant operations platform on Next.js, Supabase, Vercel, and Anthropic API.",
                "Built AI assisted workflows and clear handoffs that made daily operations easier to run.",
            ],
        ),
        (
            "Technical Product Manager | Ingram Micro",
            "Nov 2023 Dec 2024",
            [
                "Led a 60 country rollout of a unified invoice platform across SAP S/4HANA, SAP BRIM, Salesforce, and adjacent billing systems.",
                "Deployed OpenAI API and Python workflow automation across 20+ countries on the Xvantage platform.",
            ],
        ),
        (
            "Product Owner / Associate Integration Lead | Ingram Micro",
            "Nov 2022 Nov 2023",
            [
                "Owned roadmap alignment and migration coordination across SAP and cloud platform initiatives.",
                "Translated business goals into user stories, acceptance criteria, and sprint ready delivery plans.",
            ],
        ),
        (
            "Application Support Supervisor | loanDepot",
            "Apr 2019 Nov 2022",
            [
                "Reduced a 1,000+ ticket backlog to zero by redesigning support workflows and introducing sprint based resolution cycles.",
                "Maintained above 95% SLA compliance while coordinating engineering, InfoSec, and vendor escalations.",
            ],
        ),
    ]
    for title, date_text, bullets in roles:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, before=4, after=0, line=1.0)
        add_run(p, title, size=10.6, bold=True, color=INK)
        add_run(p, f"    {date_text}", size=8.75, color=SOFT)
        add_bullets(doc, bullets, size=9.72)

    add_heading(doc, "Education and Certifications")
    add_bullets(
        doc,
        [
            "M.S., Information Systems, University of Phoenix, 2022 2024",
            "B.S., Business Management, University of Phoenix, 2017 2020",
            "SAFe Scaled Agile Framework, Certified Scrum Master, CPMAI complete",
        ],
        size=9.72,
    )
    return doc


def build_letter():
    doc = Document()
    set_doc_defaults(doc)
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(fp, before=0, after=0, line=1.0)
    add_run(fp, "Rory Semeah | Vercel Developer Success Engineer", size=8.5, color=SOFT)

    add_job_title_block(doc, ROLE.title, "Company directed cover letter")
    add_credential_line(doc, f"{ROLE.city} | 442-461-3093 | roryleesemeah@icloud.com | rorysemeah.com")
    add_credential_line(doc, f"{TODAY.strftime('%B %d, %Y')}")
    add_para(doc, "Vercel Hiring Team", size=10.5, bold=True, color=INK, after=0)
    add_para(doc, "Re: Developer Success Engineer", size=10.5, bold=True, color=INK, after=8)

    paragraphs = [
        "I am applying for the Developer Success Engineer role because it matches the kind of work I do best: helping teams ship, solve technical problems quickly, and turn what works into repeatable guidance others can use.",
        "At RedLantern Studios and Ingram Micro I have shipped production systems on Next.js, Vercel, Supabase, OpenAI, Anthropic, AWS, Kubernetes, Terraform, and CI/CD pipelines. I have led a 60 country rollout, deployed AI automation across 20+ countries, and built tools and workflows that reduce friction for operators and developers.",
        "What draws me to Vercel is the mix of production readiness, performance, launch stabilization, and developer experience. That is the kind of environment where I can bring technical depth, calm communication, and a bias toward clear, reusable solutions.",
        "I also care about how the work is explained. If I am helping a team, I want them to leave with a fixed problem, a better playbook, and a shorter path the next time it happens.",
        "Thank you for your time and consideration. I would welcome the chance to talk about how I can help Vercel customers ship with confidence and stay stable after launch.",
    ]
    for text in paragraphs:
        add_para(doc, text, size=10.4, after=8, line=1.18)
    add_para(doc, "Warmly,", size=10.4, after=18)
    add_para(doc, "Rory Semeah", size=10.4, bold=True, after=0)
    return doc


def build_packet():
    doc = Document()
    set_doc_defaults(doc)
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(fp, before=0, after=0, line=1.0)
    add_run(fp, "Candidate Portfolio Packet", size=8.5, color=SOFT)

    add_kicker(doc, "Candidate Portfolio Packet")
    title = doc.add_paragraph()
    set_spacing(title, before=0, after=2, line=1.0)
    add_run(title, "Rory Semeah", size=28, bold=True, color=INK)
    add_run(title, "\n", size=1, color=INK)
    add_run(title, "Developer Success Engineer | Vercel Candidate Portfolio", size=12, bold=True, color=ACCENT)
    add_contact_line(
        doc,
        [
            ROLE.city,
            "roryleesemeah@icloud.com",
            "rorysemeah.com",
            "github.com/redlanternstudios",
        ],
    )
    add_divider(doc)

    append_portfolio_sections(doc)
    return doc


def append_portfolio_sections(doc):
    add_heading(doc, "Who I Am")
    add_bullets(
        doc,
        [
            "I build systems that help teams ship without chaos.",
            "I work best where product, engineering, and operations overlap and somebody needs a clear path forward.",
            "I like direct work, useful work, and work that makes the next launch easier than the last one.",
        ],
        size=10.3,
    )

    add_heading(doc, "What I Do")
    add_bullets(
        doc,
        [
            "I use Next.js and Vercel to build and launch production web products, then stay close long enough to help steady them after release.",
            "I use OpenAI and Anthropic to cut down repeat work and make internal workflows easier to follow.",
            "I work with Supabase, AWS, Kubernetes, Terraform, REST APIs, and GraphQL when the job needs real delivery support behind the product.",
            "I write documentation and playbooks so the same question does not have to get answered three times.",
        ],
        size=10.3,
    )

    add_heading(doc, "Why Vercel")
    add_bullets(
        doc,
        [
            "This role sits close to the work I enjoy most: helping teams get live, stay live, and recover quickly when something breaks.",
            "Vercel is built around speed, clarity, and a better developer experience, which matches how I like to work.",
            "I have already used Vercel as a production platform, so I know the pace and the language of real deployments.",
        ],
        size=10.3,
    )

    add_heading(doc, "Proof Map")
    proof_rows = [
        ("Go live readiness", "Authentic Hadith shipped through Apple Developer Console, TestFlight QA, and production approval."),
        ("Scale and rollout", "Ingram Micro 60 country invoice platform rollout across SAP S/4HANA, SAP BRIM, Salesforce, and adjacent billing systems."),
        ("Automation", "OpenAI API and Python workflow automation deployed across 20+ countries on Xvantage."),
        ("Support recovery", "loanDepot support backlog reduced from 1,000+ tickets to zero through workflow redesign."),
        ("Platform delivery", "Next.js, Vercel, Supabase, Anthropic, CI/CD, AWS, Kubernetes, Terraform."),
    ]
    for left, right in proof_rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.06
        add_run(p, f"{left}: ", size=10.1, bold=True, color=ACCENT)
        add_run(p, right, size=10.1, color=INK)

    add_heading(doc, "CV Snapshot")
    add_bullets(
        doc,
        [
            "Technical Product Manager, RedLantern Studios, Jan 2025 Present.",
            "Technical Product Manager, Ingram Micro, Nov 2023 Dec 2024.",
            "Product Owner / Associate Integration Lead, Ingram Micro, Nov 2022 Nov 2023.",
            "Application Support Supervisor, loanDepot, Apr 2019 Nov 2022.",
            "M.S. Information Systems and B.S. Business Management from University of Phoenix.",
        ],
        size=10.25,
    )

    add_heading(doc, "Working Style")
    add_bullets(
        doc,
        [
            "Plain language first.",
            "Evidence over hype.",
            "Fast iteration, honest scope, and clear handoff.",
            "If a match is weak, I want the questions that actually close the gap before I write the final materials.",
        ],
        size=10.25,
    )

    add_heading(doc, "How To Use This Packet")
    add_bullets(
        doc,
        [
            "Use the resume for portal uploads and ATS screening.",
            "Use the packet when sending directly to recruiters, hiring managers, or headhunters.",
            "Use the cover letter when the application asks for a short personal note or when sending directly by email.",
        ],
        size=10.25,
    )


def build_candidate_portfolio():
    doc = build_resume()
    doc.add_page_break()
    add_kicker(doc, "Candidate Portfolio")
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=2, line=1.0)
    add_run(p, "Rory Semeah", size=26, bold=True, color=INK)
    add_run(p, "\n", size=1, color=INK)
    add_run(p, "Developer Success Engineer | Vercel Candidate Portfolio", size=11, bold=True, color=ACCENT)
    add_contact_line(
        doc,
        [
            ROLE.city,
            "roryleesemeah@icloud.com",
            "rorysemeah.com",
            "github.com/redlanternstudios",
        ],
    )
    add_divider(doc)
    append_portfolio_sections(doc)
    return doc


def save_doc(doc, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def render_docx_to_pdf(docx_path: Path, out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    env["TMPDIR"] = "/private/tmp"
    cmd = [
        "/Users/rorysemeah/.codex/plugins/cache/openai-primary-runtime/documents/26.715.12143/skills/documents/render_docx.py",
        str(docx_path),
        "--output_dir",
        str(out_dir),
        "--emit_pdf",
    ]
    subprocess.run(["python3", *cmd[1:]], check=True, env=env)


def clear_previous_outputs():
    if OUTPUT.exists():
        for item in OUTPUT.iterdir():
            if item.name == ".gitkeep":
                continue
            if item.is_dir():
                shutil.rmtree(item)
            else:
                item.unlink()
    OUTPUT.mkdir(parents=True, exist_ok=True)
    RENDER_ROOT.mkdir(parents=True, exist_ok=True)
    for item in Path("/private/tmp").glob("vercel_*_check*"):
        if item.is_dir():
            shutil.rmtree(item)


def main():
    clear_previous_outputs()

    resume_doc = build_resume()
    letter_doc = build_letter()
    packet_doc = build_packet()
    candidate_portfolio_doc = build_candidate_portfolio()

    resume_docx = OUTPUT / "Rory_Semeah_Vercel_Developer_Success_Engineer_Resume.docx"
    letter_docx = OUTPUT / "Rory_Semeah_Vercel_Developer_Success_Engineer_Cover_Letter.docx"
    packet_docx = OUTPUT / "Rory_Semeah_Vercel_Developer_Success_Engineer_Candidate_Portfolio_Packet.docx"
    combined_docx = OUTPUT / "Rory_Semeah_Vercel_Developer_Success_Engineer_Candidate_Portfolio.docx"

    save_doc(resume_doc, resume_docx)
    save_doc(letter_doc, letter_docx)
    save_doc(packet_doc, packet_docx)
    save_doc(candidate_portfolio_doc, combined_docx)

    print(resume_docx)
    print(letter_docx)
    print(packet_docx)
    print(combined_docx)


if __name__ == "__main__":
    main()
